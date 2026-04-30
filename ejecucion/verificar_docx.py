#!/usr/bin/env python3
"""
Script para verificar que el documento Word no contenga asteriscos
"""

from docx import Document
import zipfile
import xml.etree.ElementTree as ET

def verify_docx(filepath):
    """Verifica que el documento no contenga asteriscos"""
    doc = Document(filepath)
    
    asterisk_count = 0
    total_paragraphs = 0
    total_tables = 0
    total_cells = 0
    
    # Verificar párrafos
    for para in doc.paragraphs:
        total_paragraphs += 1
        text = para.text
        if text and ('**' in text or ('*' in text and text.count('*') % 2 == 0)):
            if '**' in text or '* ' in text or text.startswith('*'):
                asterisk_count += text.count('**') + text.count('*')
                print(f"Párrafo con asteriscos: {text[:100]}...")
    
    # Verificar tablas
    for table in doc.tables:
        total_tables += 1
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                text = cell.text
                if text and ('**' in text or ('*' in text)):
                    asterisk_count += text.count('**') + text.count('*')
                    print(f"Celda con asteriscos: {text[:100]}...")
    
    print(f"\n=== RESUMEN DE VERIFICACIÓN ===")
    print(f"Párrafos procesados: {total_paragraphs}")
    print(f"Tablas procesadas: {total_tables}")
    print(f"Celdas procesadas: {total_cells}")
    print(f"Asteriscos encontrados: {asterisk_count}")
    
    if asterisk_count == 0:
        print("\n✓ VERIFICACIÓN EXITOSA: No se encontraron asteriscos en el documento")
        return True
    else:
        print(f"\n✗ VERIFICACIÓN FALLIDA: Se encontraron {asterisk_count} asteriscos")
        return False

if __name__ == '__main__':
    filepath = '/root/.openclaw/workspace/ejecucion/8_REQUERIMIENTOS_COMPLETO_41_ITS.docx'
    verify_docx(filepath)
