#!/usr/bin/env python3
"""
Script para generar documento Word con las 41 fichas ITS (CC-01 a CC-41)
Formato: Fuente Arial 11, sin asteriscos, tablas completas
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def clean_text(text):
    """Elimina asteriscos y caracteres markdown problemáticos"""
    # Eliminar asteriscos dobles (negritas markdown)
    text = re.sub(r'\*\*', '', text)
    # Eliminar asteriscos simples (cursiva markdown)
    text = re.sub(r'\*', '', text)
    # Limpiar espacios múltiples
    text = re.sub(r'  +', ' ', text)
    # Limpiar espacios al inicio y final
    text = text.strip()
    return text

def parse_table_lines(lines, start_idx):
    """Parsea una tabla markdown y retorna las filas y el índice final"""
    rows = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        # Si la línea está vacía o no es parte de la tabla
        if not line:
            i += 1
            continue
        # Si no es línea de tabla
        if not line.startswith('|'):
            break
        # Saltar la línea separadora (|---|)
        if line.replace('|', '').replace('-', '').replace(' ', '') == '':
            i += 1
            continue
        # Procesar fila de tabla
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    
    return rows, i

def add_formatted_table(doc, rows):
    """Agrega una tabla formateada al documento"""
    if not rows or len(rows) < 1:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = clean_text(cell_text)
                # Aplicar fuente Arial 11 a todas las celdas
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio después de tabla

def process_ficha(doc, filepath, ficha_num):
    """Procesa una ficha y agrega al documento"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Detectar inicio de tabla
        if stripped.startswith('|') and not in_table:
            in_table = True
            table_rows, i = parse_table_lines(lines, i)
            add_formatted_table(doc, table_rows)
            in_table = False
            continue
        
        # Procesar encabezados
        if stripped.startswith('# ') and not stripped.startswith('# 8.'):
            # Título principal de ficha
            text = clean_text(stripped[2:])
            if text:
                p = doc.add_heading(text, level=1)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
        elif stripped.startswith('## '):
            # Subsección 8.XX.X
            text = clean_text(stripped[3:])
            if text:
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.font.bold = True
        elif stripped.startswith('### '):
            # Sub-subsección
            text = clean_text(stripped[4:])
            if text:
                p = doc.add_heading(text, level=3)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
        elif stripped.startswith('#### '):
            # Sub-sub-subsección
            text = clean_text(stripped[5:])
            if text:
                p = doc.add_heading(text, level=4)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
        elif stripped.startswith('---'):
            # Separador horizontal - ignorar
            pass
        elif stripped:
            # Párrafo normal
            text = clean_text(stripped)
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        
        i += 1
    
    # Agregar salto de página entre fichas (excepto la última)
    doc.add_page_break()

def main():
    doc = Document()
    
    # Configurar estilos por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título del documento
    title = doc.add_heading('8. REQUERIMIENTOS ITS - 41 COMPONENTES DE CAMPO', level=0)
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
    
    doc.add_paragraph()
    
    # Procesar cada ficha
    base_path = '/root/.openclaw/workspace/ejecucion'
    
    for num in range(1, 42):
        ficha_code = f'CC-{num:02d}'
        filepath = os.path.join(base_path, f'{ficha_code}.md')
        
        if os.path.exists(filepath):
            print(f'Procesando {ficha_code}...')
            process_ficha(doc, filepath, num)
        else:
            print(f'ADVERTENCIA: No se encontró {filepath}')
    
    # Guardar documento
    output_path = '/root/.openclaw/workspace/ejecucion/8_REQUERIMIENTOS_COMPLETO_41_ITS.docx'
    doc.save(output_path)
    print(f'\nDocumento guardado en: {output_path}')
    
    # Verificar tamaño
    size = os.path.getsize(output_path)
    print(f'Tamaño: {size / 1024:.1f} KB')

if __name__ == '__main__':
    main()
